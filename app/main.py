"""
SkillSync Main Dashboard - Streamlit application entry point
"""
import streamlit as st
import sys
from pathlib import Path
import importlib

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

# Force reload of modules to avoid caching issues
try:
    if 'modules.skill_extractor' in sys.modules:
        importlib.reload(sys.modules['modules.skill_extractor'])
    if 'modules.resume_parser' in sys.modules:
        importlib.reload(sys.modules['modules.resume_parser'])
except Exception as e:
    print(f"Module reload warning: {e}")

try:
    from pipeline import SkillSyncPipeline
except Exception as e:
    st.error(f"Failed to load pipeline: {e}")
    st.stop()

# Page configuration
st.set_page_config(
    page_title="SkillSync",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for high-contrast, readable theme
st.markdown("""
<style>
    :root {
        --bg: #f7f8fb;
        --panel: #ffffff;
        --text: #0f172a;
        --muted: #475569;
        --accent: #0b5cab;
        --border: #cbd5e1;
        --sidebar: #0b1220;
        --sidebar-text: #e2e8f0;
    }
    [data-testid="stAppViewContainer"] {
        background: var(--bg);
        color: var(--text);
    }
    [data-testid="stSidebar"] {
        background: var(--sidebar);
    }
    [data-testid="stSidebar"] * {
        color: var(--sidebar-text) !important;
    }
    h1, h2, h3 {
        color: var(--accent);
    }
    .metric-card {
        background: var(--panel);
        border-left: 4px solid var(--accent);
        padding: 15px;
        border-radius: 6px;
        margin: 10px 0;
        color: var(--text);
    }
    .stTextInput input,
    .stTextArea textarea,
    .stSelectbox div[data-baseweb="select"] > div,
    .stMultiSelect div[data-baseweb="select"] > div {
        background: var(--panel) !important;
        color: var(--text) !important;
        border-color: var(--border) !important;
    }
    label, .stMarkdown, .stText, .stCaption {
        color: var(--text) !important;
    }
    .stCaption, .stMarkdown small {
        color: var(--muted) !important;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None
if 'pipeline' not in st.session_state:
    st.session_state.pipeline = SkillSyncPipeline()
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = None
if 'job_description' not in st.session_state:
    st.session_state.job_description = None

# Sidebar navigation
st.sidebar.title("🚀 SkillSync")
st.sidebar.markdown("---")

# Main page selection
page = st.sidebar.radio(
    "Navigation",
    ["📊 Dashboard", "📄 Resume Analysis", "💼 Job Matching", 
     "🎯 Skill Gap", "📚 Recommendations", "🔮 Career Prediction"],
    label_visibility="collapsed"
)

st.sidebar.markdown("---")
st.sidebar.markdown("### About")
st.sidebar.info(
    """
    **SkillSync** helps you:
    - Analyze your resume
    - Match jobs
    - Identify skill gaps
    - Get learning recommendations
    - Predict career opportunities
    """
)

# Main content area
if page == "📊 Dashboard":
    st.title("🚀 SkillSync - Your Career Growth Platform")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric(
            label="Resume Analysis",
            value="Ready",
            delta="Start →"
        )
    
    with col2:
        st.metric(
            label="Job Matching",
            value="Enabled",
            delta="Features →"
        )
    
    with col3:
        st.metric(
            label="Recommendations",
            value="Active",
            delta="Learning paths →"
        )
    
    st.markdown("---")
    
    st.header("Welcome to SkillSync")
    st.write("""
    SkillSync is your intelligent career development platform that helps you:
    
    1. **📄 Analyze Your Resume** - Extract and organize your skills and experience
    2. **💼 Match Jobs** - Compare your resume against job descriptions with ATS scoring
    3. **🎯 Identify Gaps** - See what skills you're missing for your target role
    4. **📚 Get Recommendations** - Discover courses and resources to bridge skill gaps
    5. **🔮 Predict Career Paths** - Understand potential career trajectories based on your profile
    
    ### Getting Started
    
    1. Go to **Resume Analysis** and upload your resume
    2. Paste a job description in **Job Matching**
    3. Review your **Skill Gap** analysis
    4. Check **Recommendations** for learning resources
    5. Explore **Career Prediction** for growth opportunities
    
    ---
    """)
    
    # Analysis Status
    if st.session_state.analysis_results:
        st.success("✅ Analysis completed successfully!")
        st.json(st.session_state.analysis_results)
    else:
        st.info("👉 Start by uploading your resume on the 'Resume Analysis' page")

elif page == "📄 Resume Analysis":
    st.title("📄 Resume Analysis")
    
    uploaded_file = st.file_uploader(
        "Upload your resume (PDF, DOCX, or TXT)",
        type=["pdf", "docx", "txt"]
    )
    
    if uploaded_file is not None:
        try:
            from modules.resume_parser import ResumeParser
            
            # Read file content based on type
            if uploaded_file.type == "text/plain":
                st.session_state.resume_text = uploaded_file.read().decode('utf-8')
            elif uploaded_file.type == "application/pdf":
                # Extract text from PDF with fallback method
                try:
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(uploaded_file)
                    text = ""
                    for page in pdf_reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
                    
                    # If extraction failed, try pdfplumber as fallback
                    if not text or len(text.strip()) < 50:
                        try:
                            import pdfplumber
                            pdf_reader = pdfplumber.open(uploaded_file)
                            text = ""
                            for page in pdf_reader.pages:
                                text += page.extract_text() or ""
                            pdf_reader.close()
                        except Exception:
                            pass
                    
                    st.session_state.resume_text = text if text else "No text extracted from PDF"
                except Exception as e:
                    st.warning(f"PDF extraction issue: {e}. Using alternate method.")
                    st.session_state.resume_text = "Unable to extract text"
                    
            elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                # Extract text from DOCX
                import docx
                doc = docx.Document(uploaded_file)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                st.session_state.resume_text = text
            else:
                st.error(f"Unsupported file type: {uploaded_file.type}")
                st.stop()
            
            st.success("✅ Resume uploaded successfully!")
            
            with st.expander("📋 Resume Content Preview"):
                st.text_area(
                    "Resume Text",
                    value=st.session_state.resume_text[:500] + "..." 
                    if len(st.session_state.resume_text) > 500 
                    else st.session_state.resume_text,
                    disabled=True,
                    height=200
                )
            
            # Extract information
            try:
                import importlib
                import sys
                # Force reload modules to ensure latest code is used
                if 'modules.resume_parser' in sys.modules:
                    importlib.reload(sys.modules['modules.resume_parser'])
                from modules.resume_parser import ResumeParser
                
                parser = ResumeParser()
                if not hasattr(parser, 'parse'):
                    st.error("❌ ResumeParser.parse method not found. Debug info: " + str(dir(parser)))
                else:
                    parsed_resume = parser.parse(st.session_state.resume_text)
                    
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.metric("Name", parsed_resume.get('name', 'N/A'))
                    
                    with col2:
                        st.metric("Title", parsed_resume.get('title', 'N/A'))
                    
                    with col3:
                        st.metric("Experience", f"{parsed_resume.get('years_experience', 0)} years")
                    
                    # Skills extraction
                    st.subheader("🎯 Extracted Skills")
                    
                    try:
                        # Validate text before extraction
                        if not st.session_state.resume_text or len(st.session_state.resume_text.strip()) < 20:
                            st.warning("⚠️ Resume text is too short or empty. Please upload a valid resume with content.")
                        else:
                            # Force fresh import
                            if 'modules.skill_extractor' in sys.modules:
                                del sys.modules['modules.skill_extractor']
                            
                            from modules.skill_extractor import SkillExtractor
                            extractor = SkillExtractor()
                            
                            # Verify method exists
                            if not hasattr(extractor, 'extract'):
                                st.error("❌ extract method not found. Available methods: " + str([m for m in dir(extractor) if not m.startswith('_')]))
                            else:
                                skills = extractor.extract(st.session_state.resume_text)
                                
                                if skills:
                                    col1, col2 = st.columns([2, 1])
                                    with col1:
                                        st.write(f"✅ Found {len(skills)} skills:")
                                        skill_text = ", ".join([str(s) for s in skills[:20]])
                                        if len(skills) > 20:
                                            skill_text += f", +{len(skills) - 20} more"
                                        st.success(skill_text)
                                else:
                                    st.warning("⚠️ No skills detected in resume. Make sure your resume contains skill keywords like: Python, Java, AWS, Docker, SQL, etc.")
                    except Exception as skill_error:
                        import traceback
                        st.error(f"⚠️ Error extracting skills: {str(skill_error)}")
                        st.error(f"Traceback: {traceback.format_exc()}")
            except Exception as e:
                st.error(f"❌ Error processing file: {str(e)}\nType: {type(e).__name__}")
            
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
    else:
        st.info("📤 Upload a resume file to get started")

elif page == "💼 Job Matching":
    st.title("💼 Job Matching")
    
    st.write("Paste the job description to compare against your resume")
    
    job_desc = st.text_area(
        "Job Description",
        value=st.session_state.job_description or "",
        height=300,
        placeholder="Paste the job description here...",
        key="job_desc_input"
    )
    
    # Auto-save job description
    if job_desc:
        st.session_state.job_description = job_desc
    
    if job_desc and st.session_state.resume_text:
        if st.button("Calculate ATS Score", use_container_width=True):
            try:
                from modules.ats_scorer import ATSScorer
                
                scorer = ATSScorer()
                score = scorer.score(st.session_state.resume_text, job_desc)
                details = scorer.get_details()
                
                # Display score
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("ATS Score", f"{score:.1f}%")
                
                with col2:
                    st.metric("Match Status", "Good" if score >= 70 else "Fair" if score >= 50 else "Needs Work")
                
                with col3:
                    st.metric("Recommendation", "Apply" if score >= 70 else "Prepare" if score >= 50 else "Skip")
                
                # Details
                st.subheader("Score Breakdown")
                st.write(details)
                # Run full analysis pipeline and store results for Recommendations and Career Prediction
                try:
                    analysis = st.session_state.pipeline.analyze(
                        st.session_state.resume_text,
                        job_desc
                    )
                    st.session_state.analysis_results = analysis
                    st.success("✅ Full analysis completed and recommendations generated")
                except Exception as e:
                    st.error(f"Error running full analysis pipeline: {e}")
                
            except Exception as e:
                st.error(f"Error calculating ATS score: {str(e)}")
        
        # Show quick skill preview
        st.subheader("📊 Quick Preview")
        try:
            from modules.skill_extractor import SkillExtractor
            extractor = SkillExtractor()
            resume_skills = set(extractor.extract(st.session_state.resume_text))
            job_skills = set(extractor.extract(job_desc))
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Your Skills", len(resume_skills))
            with col2:
                st.metric("Job Skills", len(job_skills))
            with col3:
                matched = len(resume_skills & job_skills)
                st.metric("Matched", matched)
        except:
            pass
    else:
        if not st.session_state.resume_text:
            st.warning("⚠️ Please upload a resume first on the 'Resume Analysis' page")
        else:
            st.info("📝 Paste a job description to get started")

elif page == "🎯 Skill Gap":
    st.title("🎯 Skill Gap Analysis")
    
    if st.session_state.resume_text and st.session_state.job_description:
        try:
            from modules.skill_extractor import SkillExtractor
            from modules.gap_analyzer import GapAnalyzer
            
            extractor = SkillExtractor()
            resume_skills = set(extractor.extract(st.session_state.resume_text))
            job_skills = set(extractor.extract(st.session_state.job_description))
            
            analyzer = GapAnalyzer()
            gap_analysis = analyzer.analyze(resume_skills, job_skills)
            
            matched = gap_analysis.get('matched', [])
            missing = gap_analysis.get('missing', [])
            excess = gap_analysis.get('excess', [])
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("Matched Skills", len(matched))
            
            with col2:
                st.metric("Missing Skills", len(missing))
            
            with col3:
                match_pct = int((len(matched) / max(len(job_skills), 1)) * 100) if job_skills else 0
                st.metric("Match %", f"{match_pct}%")
            
            st.markdown("---")
            
            # Matched skills
            if matched:
                st.subheader("✅ Skills You Have")
                st.success(", ".join(sorted(matched)[:15]))
                if len(matched) > 15:
                    st.caption(f"...and {len(matched) - 15} more")
            
            # Missing skills
            if missing:
                st.subheader("❌ Skills You Need to Learn")
                # Extract skill names from dicts if needed
                missing_skills_names = [m['skill'] if isinstance(m, dict) else m for m in missing]
                st.warning(", ".join(sorted(missing_skills_names)[:15]))
                if len(missing) > 15:
                    st.caption(f"...and {len(missing) - 15} more")
            
            # Excess skills
            if excess:
                st.subheader("➕ Extra Skills You Have")
                st.info(", ".join(sorted(excess)[:15]))
                if len(excess) > 15:
                    st.caption(f"...and {len(excess) - 15} more")
            
            if not matched and not missing and not excess:
                st.info("📊 Unable to extract skills from the provided texts. Try uploading a different resume or job description.")
        except Exception as e:
            st.error(f"Error analyzing skill gap: {str(e)}")
    else:
        st.warning("⚠️ Complete the following steps first:")
        col1, col2 = st.columns(2)
        with col1:
            if not st.session_state.resume_text:
                st.error("1. Upload a resume (📄 Resume Analysis page)")
            else:
                st.success("1. ✅ Resume uploaded")
        with col2:
            if not st.session_state.job_description:
                st.error("2. Paste a job description (💼 Job Matching page)")
            else:
                st.success("2. ✅ Job description added")

elif page == "📚 Recommendations":
    st.title("📚 Learning Recommendations")
    
    if st.session_state.resume_text and st.session_state.job_description:
        if st.button("Regenerate recommendations", use_container_width=True):
            with st.spinner("Rebuilding recommendations..."):
                try:
                    analysis = st.session_state.pipeline.analyze(
                        st.session_state.resume_text,
                        st.session_state.job_description
                    )
                    st.session_state.analysis_results = analysis
                    st.success("✅ Recommendations updated")
                except Exception as e:
                    st.error(f"Error regenerating recommendations: {e}")
    else:
        st.info("📝 Upload a resume and add a job description first")

    if st.session_state.analysis_results and 'recommendations' in st.session_state.analysis_results:
        recommendations = st.session_state.analysis_results['recommendations']
        
        # Filters
        col1, col2 = st.columns(2)
        
        with col1:
            platform_filter = st.selectbox(
                "Filter by Platform",
                ["All"] + sorted(set(r.get('platform', 'Unknown') for r in recommendations))
            )
        
        with col2:
            cost_filter = st.selectbox(
                "Filter by Cost",
                ["All", "Free", "Paid"]
            )

        show_all = st.checkbox("Show all recommendations (ignore filters)", value=False)
        
        def is_free_course(course: dict) -> bool:
            if 'is_free' in course:
                return bool(course.get('is_free'))
            try:
                return float(course.get('price', 0)) == 0
            except (TypeError, ValueError):
                return False
        
        filtered = recommendations
        if not show_all:
            if platform_filter != "All":
                filtered = [r for r in filtered if r.get('platform', 'Unknown') == platform_filter]
            if cost_filter == "Free":
                filtered = [r for r in filtered if is_free_course(r)]
            elif cost_filter == "Paid":
                filtered = [r for r in filtered if not is_free_course(r)]
        
        st.caption(f"Showing {len(filtered)} of {len(recommendations)} recommendations")
        
        if not filtered:
            st.info("No recommendations match those filters. Try a different platform or cost option.")
        
        # Display recommendations
        for rec in filtered[:10]:
            title = rec.get('title') or rec.get('course_name') or 'Course'
            platform = rec.get('platform', 'Unknown')
            with st.expander(f"📖 {title} - {platform}"):
                price = rec.get('price', 'N/A')
                price_text = "Free" if is_free_course(rec) else price
                st.write(f"**Description:** {rec.get('description', 'N/A')}")
                st.write(f"**Duration:** {rec.get('duration', 'N/A')}")
                st.write(f"**Level:** {rec.get('level', 'N/A')}")
                st.write(f"**Rating:** {rec.get('rating', 'N/A')} ⭐")
                st.write(f"**Price:** {price_text}")
                skills = rec.get('skills', [])
                if skills:
                    st.write(f"**Skills:** {', '.join(skills[:6])}")
                if rec.get('url'):
                    st.markdown(f"[Open course]({rec['url']})")
    else:
        st.info("📝 Run a full analysis first to get recommendations")

elif page == "🔮 Career Prediction":
    st.title("🔮 Career Prediction")
    
    if st.session_state.analysis_results and 'career_predictions' in st.session_state.analysis_results:
        careers = st.session_state.analysis_results['career_predictions']
        
        st.subheader("Predicted Career Paths")
        
        for career in careers[:5]:
            with st.expander(f"💼 {career.get('title', 'Career')} - {career.get('confidence', 0):.0f}%"):
                st.write(f"**Confidence:** {career.get('confidence', 0):.1f}%")
                st.write(f"**Description:** {career.get('description', 'N/A')}")
                st.write(f"**Required Skills:** {', '.join(career.get('required_skills', [])[:5])}")
                st.write(f"**Typical Salary:** {career.get('salary_range', 'N/A')}")
    else:
        st.info("📝 Run a full analysis first to see career predictions")

# Footer
st.markdown("---")
st.markdown(
    """
    <div style="text-align: center; color: #888; font-size: 12px;">
    SkillSync © 2024 | Powered by AI | All rights reserved
    </div>
    """,
    unsafe_allow_html=True
)
